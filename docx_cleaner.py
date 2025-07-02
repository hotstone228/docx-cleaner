from docx import Document
from docx.oxml import OxmlElement
import re
import glob


def clear_headers_and_footers(doc):
    """
    Completely clear all headers and footers from all sections by removing their XML parts.

    Args:
        doc (Document): The python-docx Document object to modify.
    """
    # Iterate through each section
    for section in doc.sections:
        # Access the section's XML element
        sectPr = section._sectPr

        # Remove header references
        for header_ref in sectPr.xpath(".//w:headerReference"):
            sectPr.remove(header_ref)

        # Remove footer references
        for footer_ref in sectPr.xpath(".//w:footerReference"):
            sectPr.remove(footer_ref)

        # Clear any existing header/footer content (precautionary)
        section.header.paragraphs.clear()
        section.footer.paragraphs.clear()
        if section.different_first_page_header_footer:
            section.first_page_header.paragraphs.clear()
            section.first_page_footer.paragraphs.clear()
        if not section.header.is_linked_to_previous:
            section.even_page_header.paragraphs.clear()
            section.even_page_footer.paragraphs.clear()


def remove_words_from_docx(input_path, output_path, words_to_remove):
    """
    Remove specified words and all headers/footers from a DOCX file and save the result.

    Args:
        input_path (str): Path to the input DOCX file.
        output_path (str): Path where the modified DOCX file will be saved.
        words_to_remove (list): List of words to erase from the document.
    """
    # Load the DOCX file
    doc = Document(input_path)

    # Remove all headers and footers
    clear_headers_and_footers(doc)

    # Process each paragraph in the document (main body only)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for word in words_to_remove:
                # Remove whole words only
                pattern = re.escape(word)
                text = re.sub(pattern, "", text)
            # Clean up spaces
            text = re.sub(r"\s+", " ", text)
            text = re.sub(r"\s+([.,!?])", r"\1", text)
            run.text = text

    # Save the modified document
    doc.save(output_path)


# Example usage
if __name__ == "__main__":
    # Load the list of words to remove from a file
    with open("words_to_remove.txt") as input_file:
        words_to_remove = set(line.strip() for line in input_file)

    # Specify input and output file paths
    input_files = glob.glob("*.doc*")
    for input_file in input_files:
        output_file = "new_" + input_file
        # Call the function to process the document
        remove_words_from_docx(input_file, output_file, words_to_remove)
        print(
            f"Processed '{input_file}' and saved as '{output_file}'. Headers and footers removed."
        )
